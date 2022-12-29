from pdf2docx import parse
from os import scandir
import docx2txt
from docx import Document
import pandas as pd
import warnings
warnings.filterwarnings("ignore")
from nltk import tokenize
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import numpy as np
import fitz
from summarizer import Summarizer
from rouge_score import rouge_scorer
import torch
from summarizer import Summarizer
from summarizer import TransformerSummarizer
import transformers
import re
# from __future__ import (
#     absolute_import, division, print_function, unicode_literals
# )

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

class TextSumerizationPDF():
    def __init__(self,name_pdf):

        self.pdf_file = name_pdf


        self.total_text = self.read_only_pdf_text()
        self.text_without_abstract, self.orginal_abstract = self.remove_abstract_from_pdf()
        self.text_after_reprocessing = self.pdf_preprocessing()
        self.len_abstract = len(self.orginal_abstract.split('.'))
        self.all_blocks = self.get_blocks_from_pdf()
        self.df_pages = self.get_data_from_blocks()
        # -------------------------------------------------------------------models and result
        self.model_abstract = self.bert_extractive_summarizer()
        print("bert_extractive_summarizer: " + self.model_abstract)

        # get_data_of_result help function in line 
        self.tableDF = self.get_data_of_result('bert_results')
        print(self.tableDF)

        self.model_abstract = self.GPT2()
        print("GPT2: " + self.model_abstract)

        # get_data_of_result help function in line 
        self.tableDF = self.get_data_of_result('GPT2_results')
        print(self.tableDF)

    def iter_block_items(self, parent):
        """
        Generate a reference to each paragraph and table child within *parent*,
        in document order. Each returned value is an instance of either Table or
        Paragraph. *parent* would most commonly be a reference to a main
        Document object, but also works for a _Cell object, which itself can
        contain paragraphs and tables.
        """
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
            # print(parent_elm.xml)
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)


    def read_only_pdf_text(self):

        word_file = "text_word.docx"

        # A wrapped method parse() to convert all/specified pdf pages to docx. 
        # Multi-processing is supported in case pdf file with a large number of pages.
        parse(self.pdf_file, word_file, start=0, end=None)

        # process extract text and lower() method 
        # returns a string where all characters are lower case.
        pdf_all_data = docx2txt.process(word_file).lower()
        
        # remove all text after references words
        text_without_references = pdf_all_data.split('references', 1)[0]

        # Return a |Document| object loaded from *docx*, where *docx* can be
        # either a path to a ``.docx`` file (a string) or a file-like object. If
        # *docx* is missing or ``None``, the built-in default document "template"
        # is loaded.
        document = Document()
        
        # This method returns a reference to a paragraph, newly added paragraph
        #  at the end of the document. In this case all text without references.
        document.add_paragraph(text_without_references)

        # Save new documents
        document.save('only_text.docx')

        total_text = []

        # remove tables and images
        document = Document('only_text.docx')

        # iter_block_items help function in line 54
        for item in self.iter_block_items(document):
            total_text.append(item.text if isinstance(item, Paragraph) else '<table>')
        return "".join(total_text)


    def remove_abstract_from_pdf(self):
        split_page = self.total_text.split('.')
        abstract_sent=[sentence + '.' for sentence in split_page if 'abstract' in sentence][0]
        abstract_sent= abstract_sent[abstract_sent.find('abstract'):]
        keywords_sent=[sentence + '.' for sentence in split_page if 'keywords' in sentence][0]
        keywords_sent= keywords_sent[keywords_sent.find('keywords'):]
        data_batween = self.total_text[self.total_text.rfind(abstract_sent)+len(abstract_sent):self.total_text.rfind(keywords_sent)]
        orginal_abstract = abstract_sent + data_batween + keywords_sent
        text_without_abstract= self.total_text.replace(orginal_abstract, '')
        return  text_without_abstract, orginal_abstract

    def data_preprocessing(self, row):
        # Lowering letters
        row = row.str.lower()
        # remove fig and table tag
        row = row.str.replace("fig.", "")
        row = row.str.replace("<table>", "")
        # Removing html tags
        row = row.str.replace(r'<[^<>]*>', ' ', regex=True)
        # Removing urls
        row = row.str.replace(r'http\S+', ' ', regex=True)
        row = row.str.replace(r'www\S+', ' ', regex=True)
        # remove ascii
        row = row.str.replace('(([\\xbc-\\xbe])?)', '')
        # row = row.str.replace(r'[^\x00-\x7F\x80-\xFF\u0100-\u017F\u0180-\u024F\u1E00-\u1EFF]', '')
        # row = row.str.encode('ascii', 'ignore').str.decode('ascii')
        # Removing numbers
        row = row.str.replace('\d+', ' ')
        # Remove special characers
        row = row.str.replace('\W', ' ')
        # Distribution of tokens
        row = row.apply(word_tokenize)
        # Remove stopwords
        stop_words = set(stopwords.words('english'))
        row =row.apply(lambda x: [item for item in x if item not in stop_words])
        # Remove roman numbers
        roman_numbers = ['i','ii', 'iii', 'iv', 'v','vi', 'vii', 'viii']
        row =row.apply(lambda x: [item for item in x if item not in roman_numbers])
        # Convert lists to strings
        row = row.apply(lambda x: ' '.join(map(str, x)))
        return row


    def pdf_preprocessing(self):
        sentenceSplit = tokenize.sent_tokenize(self.text_without_abstract)
        df = pd.DataFrame(sentenceSplit,columns =['Sentence'])
        df["Sentence"] = self.data_preprocessing(df["Sentence"])
        df =pd.DataFrame(['. '.join(df['Sentence'].to_list())], columns=['Sentence'])
        text_after_reprocessing = df.values.tolist()[0][0]
        return text_after_reprocessing

    def evaluation(self, model_abstract):
        scorer = rouge_scorer.RougeScorer(['rouge1', 'rouge2', 'rougeL'], use_stemmer=True)
        scores = scorer.score(model_abstract, self.orginal_abstract)
        return scores

    def bert_extractive_summarizer(self):
        model = Summarizer()
        model_abstract = model(self.text_after_reprocessing, num_sentences=self.len_abstract)
        scores = self.evaluation(model_abstract)
        print(scores)
        return model_abstract

    def GPT2(self):
        GPT2_model = TransformerSummarizer(transformer_type="GPT2",transformer_model_key="gpt2-medium")
        model_abstract = ''.join(GPT2_model(self.text_after_reprocessing, num_sentences=self.len_abstract))
        scores = self.evaluation(model_abstract)
        print(scores)
        return model_abstract

    def get_blocks_from_pdf(self):
        all_blocks =[]
        for page in fitz.open(self.pdf_file):

             # “blocks”: generate a list of text blocks (= paragraphs).
            all_blocks.append(page.get_text("blocks"))
   
        return  all_blocks

    def conect_paragraph_to_each_line(self, df_pages):
        all_paragraph = []
        for list_elem in df_pages['paragraph'].values.tolist():
            val_paragraph, temp_paragraph = 0, []
            for word in list_elem:
                if (type(word) == str):
                    val_paragraph = word
                temp_paragraph.append(val_paragraph)
            all_paragraph.append(temp_paragraph)
        return all_paragraph

    def get_data_of_paragraph(self, df_blocks):
        paragraphDf = pd.DataFrame()
        df_blocks['paragraph_length'] = df_blocks['text'].str.split().apply(len)
        paragraphDf = df_blocks[df_blocks['paragraph_length'] <= 4]

        paragraphDf['text'] = paragraphDf['text'].str.lower()

        # find if there is a substring such that it has @|–|www|references
        paragraphDf = paragraphDf[~paragraphDf["text"].str.contains('(@|–|www|references)', regex=True)]

        # data_preprocessing help function
        paragraphDf["text"] = self.data_preprocessing(paragraphDf["text"])
    
        # Access group of values using labels.
        paragraphDf = paragraphDf.loc[paragraphDf["text"] != '']

        # Pandas rename() method is used to rename any index, column or row.
        paragraphDf = paragraphDf[['page', 'line', 'text']].rename(
            columns={'page': 'page_paragraph', 'line': 'line_paragraph', 'text': 'paragraph'})
        return paragraphDf

    def get_data_from_blocks(self):
        all_data_blocks = [(page_ind,) + tup for page_ind in range(len(self.all_blocks)) for tup in
                           self.all_blocks[page_ind]]
        df_blocks = pd.DataFrame(all_data_blocks,
                                 columns=['page', 'x0', 'y0', 'x1', 'y1', 'text', 'line', 'img/text_index'])


        df_blocks = df_blocks[df_blocks["img/text_index"] == 0]
        df_blocks = df_blocks[['page', 'text', 'line']]

        # data_preprocessing help function in lne
        df_blocks["text"] = self.data_preprocessing(df_blocks["text"])

        # get_data_of_paragraph help function in line
        paragraphDf = self.get_data_of_paragraph(df_blocks)

        # Access group of values using labels.
        df_blocks = df_blocks.loc[df_blocks["text"] != '']

        # Pandas concat() method is used to concatenate pandas objects such as DataFrames and Series.
        # if axis=1, then column-wise concatenation is performed
        all_data_sent = pd.concat([df_blocks, paragraphDf], axis=1)
        all_data_sent = all_data_sent[['page', 'text', 'line', 'paragraph']]

        # The agg() method allows to apply a function or a list of function names to be executed 
        # along one of the axis of the DataFrame, default 0, which is the index (row) axis.
        df_pages = all_data_sent.groupby('page').agg(
            {'text': ' '.join, 'line': lambda x: x.tolist(), 'paragraph': lambda x: x.tolist()}).reset_index()

        # conect_paragraph_to_each_line help function in line
        all_paragraph = self.conect_paragraph_to_each_line(df_pages)
        df_pages['paragraph'] = pd.DataFrame([all_paragraph]).T

        return df_pages

    def get_data_of_result(self, save_file_name):
        tableDF =pd.DataFrame()
        for sent in self.model_abstract.split('.'):
            if(len(sent)>=1):
                tableSent =pd.DataFrame([sent]).rename(columns={0:'sent'})
                dfContains= self.df_pages[self.df_pages.text.str.contains(sent)].reset_index()
                if(len(dfContains.index)>0):
                    dfTableSent =pd.concat([tableSent]*dfContains.shape[0], ignore_index=True)
                    tableDF = tableDF.append(pd.concat([dfTableSent, dfContains], axis =1, join="inner" ))
        tableDF=tableDF[['sent','page','text','line', 'paragraph']]
        tableDF['page'] = tableDF['page'] + 1
        tableDF = tableDF.loc[tableDF["sent"] !=' '].rename(columns={'text':'all text in page'})
        tableDF.to_csv(save_file_name)
        return tableDF

if __name__ == "__main__":
    name_pdf = '2.pdf'
    TextSumerizationPDF(name_pdf)